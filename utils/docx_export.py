import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_heading_color(paragraph, hex_color="2E4057"):
    """Apply a hex color to all runs in a heading paragraph."""
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _add_metadata_row(table, label, value):
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]
    label_cell.paragraphs[0].add_run(label).bold = True
    label_cell.paragraphs[0].runs[0].font.size = Pt(9)
    p = value_cell.paragraphs[0]
    p.add_run(str(value))
    p.runs[0].font.size = Pt(9)


def build_docx(
    url: str,
    page_type: str,
    template_name: str,
    primary_keyword: str,
    section_results: dict,
    template_sections: list,
    keyword_assignment: dict,
    word_count: int,
    competitor_urls: list,
    h1: str = "",
) -> bytes:
    """
    Builds a .docx file in memory and returns it as bytes.
    Structured as:
      - Metadata table at top
      - Full page copy with H1/H2/H3 headings
      - Diagnostic section at bottom
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Default font ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Metadata block ────────────────────────────────────────────────────────
    meta_heading = doc.add_paragraph("Page Metadata")
    meta_heading.style = "Heading 2"
    _set_heading_color(meta_heading, "666666")

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(5.0)

    _add_metadata_row(table, "URL", url)
    _add_metadata_row(table, "Page Type", page_type.replace("_", " ").title())
    _add_metadata_row(table, "Template", template_name)
    _add_metadata_row(table, "Primary Keyword", primary_keyword)
    _add_metadata_row(table, "Total Word Count", str(word_count))

    doc.add_paragraph()

    # ── Page copy ─────────────────────────────────────────────────────────────
    copy_heading = doc.add_paragraph("Generated Copy")
    copy_heading.style = "Heading 2"
    _set_heading_color(copy_heading, "666666")
    doc.add_paragraph()

    for section in template_sections:
        sec_name = section["name"]
        text = section_results.get(sec_name, "")
        if not text:
            continue

        lines = text.splitlines()
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                p = doc.add_paragraph(line[2:], style="Heading 1")
                _set_heading_color(p, "2E4057")
            elif line.startswith("## "):
                p = doc.add_paragraph(line[3:], style="Heading 2")
                _set_heading_color(p, "2E4057")
            elif line.startswith("### "):
                p = doc.add_paragraph(line[4:], style="Heading 3")
                _set_heading_color(p, "2E4057")
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", line):
                content = re.sub(r"^\d+\.\s", "", line)
                p = doc.add_paragraph(content, style="List Number")
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()

    # ── Diagnostic block ──────────────────────────────────────────────────────
    doc.add_page_break()
    diag_heading = doc.add_paragraph("Diagnostics")
    diag_heading.style = "Heading 2"
    _set_heading_color(diag_heading, "999999")

    if competitor_urls:
        p = doc.add_paragraph()
        p.add_run("Competitors scraped: ").bold = True
        p.add_run(", ".join(competitor_urls))
        p.runs[0].font.size = Pt(9)
        p.runs[1].font.size = Pt(9)
    else:
        p = doc.add_paragraph()
        p.add_run("Competitors scraped: ").bold = True
        p.add_run("None")
        p.runs[0].font.size = Pt(9)
        p.runs[1].font.size = Pt(9)

    diag_table = doc.add_table(rows=0, cols=4)
    diag_table.style = "Table Grid"

    # Header row
    header_row = diag_table.add_row()
    for i, label in enumerate(["Section", "Words", "Keyword Slot", "Keyword Used"]):
        cell = header_row.cells[i]
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(8)

    for section in template_sections:
        sec_name = section["name"]
        assign = keyword_assignment.get(sec_name, {})
        primary = assign.get("primary", "")
        supporting = assign.get("supporting", "")
        text = section_results.get(sec_name, "")
        wc = len(text.split()) if text else 0
        kw_slot = section.get("keyword_slot", "none")
        kw_used = primary if primary else supporting if supporting else ""

        row = diag_table.add_row()
        for i, val in enumerate([section["label"], str(wc), kw_slot, kw_used]):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(8)

    # ── Save to bytes ──────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
